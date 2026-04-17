import re
from io import BytesIO
from pathlib import Path

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from fpdf import FPDF

st.set_page_config(page_title="Audit Dashboard", layout="wide")
DATA_DIR = Path(__file__).parent / "data"

st.markdown(
    """
    <style>
    :root {
        --bg: #f3efe7;
        --panel: #fffdf8;
        --panel-alt: #f8f2e8;
        --ink: #1f2937;
        --muted: #6b7280;
        --line: #ddd3c3;
        --accent: #b45309;
        --accent-soft: #f59e0b;
        --danger: #b91c1c;
        --success: #166534;
    }

    .stApp {
        background:
            radial-gradient(circle at top right, rgba(245, 158, 11, 0.12), transparent 22%),
            linear-gradient(180deg, #f8f3eb 0%, #f2ede4 100%);
    }

    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }

    .hero {
        background: linear-gradient(135deg, rgba(24, 24, 27, 0.96), rgba(120, 53, 15, 0.92));
        color: white;
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 24px;
        padding: 30px 32px;
        margin-bottom: 1.25rem;
        box-shadow: 0 24px 60px rgba(31, 41, 55, 0.15);
    }

    .hero h1 {
        margin: 0;
        font-size: 3rem;
        line-height: 1;
        letter-spacing: -0.04em;
    }

    .hero p {
        margin: 0.75rem 0 0;
        color: rgba(255,255,255,0.82);
        font-size: 1.02rem;
        max-width: 780px;
    }

    .section-label {
        color: var(--accent);
        text-transform: uppercase;
        letter-spacing: 0.14em;
        font-size: 0.78rem;
        font-weight: 700;
        margin-bottom: 0.6rem;
    }

    .metric-card {
        background: rgba(255, 253, 248, 0.92);
        border: 1px solid var(--line);
        border-radius: 20px;
        padding: 20px 22px;
        min-height: 132px;
        box-shadow: 0 10px 30px rgba(31, 41, 55, 0.06);
    }

    .metric-label {
        color: var(--muted);
        font-size: 0.92rem;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        margin-bottom: 0.7rem;
    }

    .metric-value {
        color: var(--ink);
        font-size: 2.3rem;
        font-weight: 800;
        line-height: 1;
        margin-bottom: 0.5rem;
    }

    .metric-note {
        color: var(--muted);
        font-size: 0.95rem;
    }

    .panel {
        background: rgba(255, 253, 248, 0.88);
        border: 1px solid var(--line);
        border-radius: 22px;
        padding: 18px 18px 10px;
        margin-bottom: 1rem;
        box-shadow: 0 10px 30px rgba(31, 41, 55, 0.05);
    }

    div[data-testid="stDataFrame"] {
        border-radius: 16px;
        overflow: hidden;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def load_doc(file):
    file_bytes = file.getvalue()
    return Document(BytesIO(file_bytes))


def extract_summary_table(file):
    doc = load_doc(file)
    data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]

            if len(row_data) >= 5 and ("Failed" in row_data or "Passed" in row_data):
                try:
                    data.append(
                        {
                            "Asset": row_data[1],
                            "Vulnerability": row_data[2],
                            "Status": row_data[3],
                            "Recommendation": row_data[4],
                            "Reference": row_data[5] if len(row_data) > 5 else "",
                        }
                    )
                except Exception:
                    continue

    return pd.DataFrame(data)


def extract_summary_counts(file):
    doc = load_doc(file)
    text = "\n".join([p.text for p in doc.paragraphs])

    passed, failed = 0, 0

    passed_match = re.search(r"(\d+)\s*Passed", text)
    failed_match = re.search(r"(\d+)\s*Failed", text)

    if passed_match:
        passed = int(passed_match.group(1))
    if failed_match:
        failed = int(failed_match.group(1))

    return passed, failed


def generate_ai_recommendation(vuln):
    vuln = vuln.lower()

    if "ssh" in vuln:
        return "Enable SSH and disable Telnet."
    if "vlan" in vuln:
        return "Ensure VLAN segmentation."
    if "password" in vuln:
        return "Use SHA-512 encryption."
    if "telnet" in vuln:
        return "Disable Telnet."
    return "Follow CIS/NIST standards."


def generate_pdf(summary_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary_text)
    return pdf.output(dest="S").encode("latin-1")


def render_metric_card(label, value, note):
    st.markdown(
        f"""
        <div class="metric-card">
            <div class="metric-label">{label}</div>
            <div class="metric-value">{value}</div>
            <div class="metric-note">{note}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


@st.cache_data
def load_framework_data():
    cis_df = pd.read_csv(DATA_DIR / "cis_controls.csv")
    nist_df = pd.read_csv(DATA_DIR / "nist_controls.csv")
    mapping_df = pd.read_csv(DATA_DIR / "control_mappings.csv")
    return cis_df, nist_df, mapping_df


def classify_dependency(matches):
    frameworks = set(matches["Framework"].tolist())
    if len(matches) > 1 and len(frameworks) > 1:
        return "Cross-framework dependency"
    if len(matches) > 1:
        return "Multi-control dependency"
    if len(matches) == 1:
        return "Direct control dependency"
    return "Not mapped"


def map_finding_to_controls(vulnerability, recommendation, control_df):
    finding_text = f"{vulnerability} {recommendation}".lower()
    matches = []

    for _, row in control_df.iterrows():
        keywords = [keyword.strip().lower() for keyword in str(row["keywords"]).split(",")]
        hit_count = sum(1 for keyword in keywords if keyword and keyword in finding_text)
        if hit_count:
            confidence = min(0.35 + (0.18 * hit_count), 0.95)
            matches.append(
                {
                    "Framework": row["framework"],
                    "Control ID": row["control_id"],
                    "Control Title": row["title"],
                    "Control Category": row["category"],
                    "Dependency Tags": row["dependency_tags"],
                    "Match Confidence": round(confidence, 2),
                }
            )

    if not matches:
        return pd.DataFrame(
            [
                {
                    "Framework": "Unmapped",
                    "Control ID": "No direct benchmark mapping",
                    "Control Title": "Manual GRC review required",
                    "Control Category": "Needs review",
                    "Dependency Tags": "manual assessment",
                    "Match Confidence": 0.0,
                }
            ]
        )

    matches_df = pd.DataFrame(matches).sort_values(
        ["Match Confidence", "Framework"], ascending=[False, True]
    )
    return matches_df.drop_duplicates(subset=["Framework", "Control ID"])


def build_grc_register(findings_df, control_library, control_links):
    grc_rows = []

    for _, row in findings_df.iterrows():
        matches = map_finding_to_controls(
            row["Vulnerability"], row["Recommendation"], control_library
        )
        dependency_status = classify_dependency(matches)

        for _, match in matches.iterrows():
            related_controls = control_links[
                control_links["primary_control"].eq(match["Control ID"])
                | control_links["related_control"].eq(match["Control ID"])
            ]
            related_summary = (
                "; ".join(
                    sorted(
                        {
                            f'{link["primary_control"]} <-> {link["related_control"]}'
                            for _, link in related_controls.iterrows()
                        }
                    )
                )
                if not related_controls.empty
                else "No predefined linked controls"
            )

            grc_rows.append(
                {
                    "Asset": row["Asset"],
                    "Source File": row["Source File"],
                    "Vulnerability": row["Vulnerability"],
                    "Status": row["Status"],
                    "Compliance": row["Compliance"],
                    "Risk Level": row["Risk Level"],
                    "Framework": match["Framework"],
                    "Control ID": match["Control ID"],
                    "Control Title": match["Control Title"],
                    "Control Category": match["Control Category"],
                    "Dependency Tags": match["Dependency Tags"],
                    "Dependency Status": dependency_status,
                    "Match Confidence": match["Match Confidence"],
                    "Linked Controls": related_summary,
                }
            )

    return pd.DataFrame(grc_rows)


st.markdown(
    """
    <div class="hero">
        <div class="section-label">Cyber Audit Intelligence</div>
        <h1>Audit Dashboard</h1>
        <p>
            Transform raw audit reports into an executive view of compliance posture,
            critical failures, and remediation priorities across assets.
        </p>
    </div>
    """,
    unsafe_allow_html=True,
)

uploaded_files = st.file_uploader(
    "Upload Audit Report (.docx)", type=["docx"], accept_multiple_files=True
)

if not uploaded_files:
    st.info("Upload one or more .docx audit reports to view the dashboard.")
else:
    all_data = []
    total_passed = 0
    total_failed = 0

    for file in uploaded_files:
        df = extract_summary_table(file)
        passed, failed = extract_summary_counts(file)

        total_passed += passed
        total_failed += failed

        if not df.empty:
            df["Source File"] = file.name
            all_data.append(df)

    if not all_data:
        st.warning("No findings were detected in the uploaded reports.")
    else:
        df = pd.concat(all_data, ignore_index=True)
        cis_df, nist_df, control_mapping_df = load_framework_data()
        cis_df["framework"] = "CIS"
        nist_df["framework"] = "NIST"
        control_library = pd.concat([cis_df, nist_df], ignore_index=True)
        df["Asset"] = df["Asset"].fillna("Unknown Asset")
        df["Vulnerability"] = df["Vulnerability"].fillna("Unspecified Finding")
        df["Status"] = df["Status"].fillna("Unknown")
        df["Recommendation"] = df["Recommendation"].fillna("")
        df["Reference"] = df["Reference"].fillna("")

        df["Compliance"] = df["Status"].apply(
            lambda x: "Non-Compliant" if "Fail" in str(x) else "Compliant"
        )

        def risk_score(status):
            status = str(status)
            if "Fail" in status:
                return 9
            if "Partial" in status:
                return 5
            return 1

        df["Risk Score"] = df["Status"].apply(risk_score)

        def risk_level(score):
            if score >= 7:
                return "Critical"
            if score >= 4:
                return "Medium"
            return "Low"

        df["Risk Level"] = df["Risk Score"].apply(risk_level)
        df["AI Recommendation"] = df["Vulnerability"].apply(generate_ai_recommendation)

        total = total_passed + total_failed
        compliance = round((total_passed / total) * 100, 2) if total > 0 else 0
        critical_count = int((df["Risk Level"] == "Critical").sum())
        unique_assets = int(df["Asset"].nunique())
        top_asset_risk = (
            df.groupby("Asset", as_index=False)["Risk Score"].sum()
            .sort_values("Risk Score", ascending=False)
            .head(1)
        )
        highest_risk_asset = (
            top_asset_risk.iloc[0]["Asset"] if not top_asset_risk.empty else "N/A"
        )

        with st.sidebar:
            st.markdown("### Filters")
            selected_asset = st.multiselect(
                "Asset",
                options=sorted(df["Asset"].unique().tolist()),
                default=sorted(df["Asset"].unique().tolist()),
            )
            selected_risk = st.multiselect(
                "Risk Level",
                options=["Critical", "Medium", "Low"],
                default=["Critical", "Medium", "Low"],
            )
            selected_source = st.multiselect(
                "Source File",
                options=sorted(df["Source File"].unique().tolist()),
                default=sorted(df["Source File"].unique().tolist()),
            )

        filtered_df = df[
            df["Asset"].isin(selected_asset)
            & df["Risk Level"].isin(selected_risk)
            & df["Source File"].isin(selected_source)
        ].copy()

        if filtered_df.empty:
            st.warning("The current filters removed all findings. Adjust the sidebar filters.")
            st.stop()

        severity_order = ["Critical", "Medium", "Low"]
        filtered_df["Risk Level"] = pd.Categorical(
            filtered_df["Risk Level"], categories=severity_order, ordered=True
        )

        risk_distribution = (
            filtered_df.groupby("Risk Level", as_index=False)
            .size()
            .rename(columns={"size": "Count"})
            .sort_values("Risk Level")
        )
        asset_risk = (
            filtered_df.groupby("Asset", as_index=False)
            .agg(
                Total_Risk=("Risk Score", "sum"),
                Findings=("Vulnerability", "count"),
                Failed=("Compliance", lambda s: int((s == "Non-Compliant").sum())),
            )
            .sort_values(["Total_Risk", "Failed"], ascending=False)
        )
        top_findings = (
            filtered_df.groupby(["Vulnerability", "Risk Level"], as_index=False)
            .agg(
                Count=("Vulnerability", "count"),
                Avg_Risk=("Risk Score", "mean"),
            )
            .sort_values(["Count", "Avg_Risk"], ascending=False)
            .head(10)
        )
        recommendation_df = (
            filtered_df[filtered_df["Compliance"] == "Non-Compliant"][
                ["Asset", "Vulnerability", "Risk Level", "AI Recommendation", "Reference"]
            ]
            .drop_duplicates()
            .sort_values("Risk Level")
        )
        grc_df = build_grc_register(filtered_df, control_library, control_mapping_df)
        framework_coverage = (
            grc_df[grc_df["Framework"] != "Unmapped"]
            .groupby("Framework", as_index=False)
            .agg(
                Controls_Matched=("Control ID", "nunique"),
                Findings_Mapped=("Vulnerability", "count"),
                Avg_Confidence=("Match Confidence", "mean"),
            )
        )
        dependency_counts = (
            grc_df.groupby("Dependency Status", as_index=False)
            .size()
            .rename(columns={"size": "Count"})
            .sort_values("Count", ascending=False)
        )
        unmapped_findings = (
            grc_df[grc_df["Framework"] == "Unmapped"][
                ["Asset", "Vulnerability", "Risk Level", "Dependency Status"]
            ]
            .drop_duplicates()
        )
        control_gap_view = (
            grc_df[grc_df["Framework"] != "Unmapped"]
            .groupby(["Framework", "Control ID", "Control Title"], as_index=False)
            .agg(
                Findings_Impacted=("Vulnerability", "count"),
                Highest_Risk=("Risk Level", "max"),
                Avg_Confidence=("Match Confidence", "mean"),
            )
            .sort_values(["Findings_Impacted", "Avg_Confidence"], ascending=False)
        )

        kpi1, kpi2, kpi3, kpi4 = st.columns(4)
        with kpi1:
            render_metric_card("Compliance Rate", f"{compliance}%", "Overall pass rate")
        with kpi2:
            render_metric_card("Critical Issues", critical_count, "Immediate action items")
        with kpi3:
            render_metric_card("Assets Reviewed", unique_assets, "Distinct systems assessed")
        with kpi4:
            render_metric_card("Highest Risk Asset", highest_risk_asset, "Largest cumulative exposure")

        col_a, col_b = st.columns([1.1, 0.9])

        with col_a:
            st.markdown('<div class="panel">', unsafe_allow_html=True)
            st.markdown("##### Compliance Posture")
            fig1 = px.pie(
                filtered_df,
                names="Compliance",
                hole=0.62,
                color="Compliance",
                color_discrete_map={
                    "Compliant": "#2f855a",
                    "Non-Compliant": "#c2410c",
                },
            )
            fig1.update_traces(textposition="inside", textinfo="percent+label")
            fig1.update_layout(
                margin=dict(l=0, r=0, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                height=340,
            )
            st.plotly_chart(fig1, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        with col_b:
            st.markdown('<div class="panel">', unsafe_allow_html=True)
            st.markdown("##### Severity Distribution")
            fig2 = px.bar(
                risk_distribution,
                x="Risk Level",
                y="Count",
                color="Risk Level",
                color_discrete_map={
                    "Critical": "#b91c1c",
                    "Medium": "#d97706",
                    "Low": "#15803d",
                },
            )
            fig2.update_layout(
                showlegend=False,
                margin=dict(l=0, r=0, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                height=340,
                xaxis_title="",
            )
            st.plotly_chart(fig2, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        col_c, col_d = st.columns([1.2, 0.8])

        with col_c:
            st.markdown('<div class="panel">', unsafe_allow_html=True)
            st.markdown("##### Asset Exposure Ranking")
            fig3 = px.bar(
                asset_risk.head(10),
                x="Total_Risk",
                y="Asset",
                orientation="h",
                color="Failed",
                color_continuous_scale=["#fed7aa", "#c2410c"],
            )
            fig3.update_layout(
                margin=dict(l=0, r=0, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                height=430,
                yaxis={"categoryorder": "total ascending"},
            )
            st.plotly_chart(fig3, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        with col_d:
            st.markdown('<div class="panel">', unsafe_allow_html=True)
            st.markdown("##### Executive Snapshot")
            st.markdown(
                f"""
                **Reports loaded:** {len(uploaded_files)}  
                **Total findings:** {len(filtered_df)}  
                **Failed checks:** {int((filtered_df["Compliance"] == "Non-Compliant").sum())}  
                **Most exposed asset:** {highest_risk_asset}
                """
            )
            gauge = go.Figure(
                go.Indicator(
                    mode="gauge+number",
                    value=compliance,
                    title={"text": "Compliance Score"},
                    gauge={
                        "axis": {"range": [0, 100]},
                        "bar": {"color": "#b45309"},
                        "steps": [
                            {"range": [0, 60], "color": "#fee2e2"},
                            {"range": [60, 85], "color": "#fef3c7"},
                            {"range": [85, 100], "color": "#dcfce7"},
                        ],
                    },
                )
            )
            gauge.update_layout(
                height=260,
                margin=dict(l=20, r=20, t=40, b=0),
                paper_bgcolor="rgba(0,0,0,0)",
            )
            st.plotly_chart(gauge, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        tab1, tab2, tab3 = st.tabs(
            ["Priority Findings", "Detailed Register", "Action Plan"]
        )

        with tab1:
            st.markdown("##### Top Repeated Vulnerabilities")
            st.dataframe(top_findings, use_container_width=True, hide_index=True)

        with tab2:
            st.markdown("##### Full Findings Register")
            st.dataframe(
                filtered_df.sort_values(["Risk Score", "Asset"], ascending=[False, True]),
                use_container_width=True,
                hide_index=True,
            )

        with tab3:
            st.markdown("##### Recommended Remediation Queue")
            st.dataframe(
                recommendation_df,
                use_container_width=True,
                hide_index=True,
            )

        st.markdown("### GRC Benchmark Analysis")
        grc_col1, grc_col2 = st.columns([0.95, 1.05])

        with grc_col1:
            st.markdown('<div class="panel">', unsafe_allow_html=True)
            st.markdown("##### Framework Coverage")
            if framework_coverage.empty:
                st.info("No built-in CIS/NIST controls were matched for the current filters.")
            else:
                fig4 = px.bar(
                    framework_coverage,
                    x="Framework",
                    y="Findings_Mapped",
                    color="Controls_Matched",
                    color_continuous_scale=["#fde68a", "#b45309"],
                    text_auto=True,
                )
                fig4.update_layout(
                    margin=dict(l=0, r=0, t=10, b=10),
                    paper_bgcolor="rgba(0,0,0,0)",
                    plot_bgcolor="rgba(0,0,0,0)",
                    height=320,
                )
                st.plotly_chart(fig4, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        with grc_col2:
            st.markdown('<div class="panel">', unsafe_allow_html=True)
            st.markdown("##### Dependency Analysis")
            fig5 = px.bar(
                dependency_counts,
                x="Dependency Status",
                y="Count",
                color="Dependency Status",
                color_discrete_map={
                    "Cross-framework dependency": "#7c2d12",
                    "Multi-control dependency": "#c2410c",
                    "Direct control dependency": "#d97706",
                    "Not mapped": "#9ca3af",
                },
            )
            fig5.update_layout(
                showlegend=False,
                margin=dict(l=0, r=0, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                height=320,
                xaxis_title="",
            )
            st.plotly_chart(fig5, use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        grc_tab1, grc_tab2, grc_tab3 = st.tabs(
            ["Control Register", "Gap View", "Unmapped Findings"]
        )

        with grc_tab1:
            st.markdown("##### Mapped CIS / NIST Controls")
            st.dataframe(
                grc_df.sort_values(
                    ["Framework", "Match Confidence"], ascending=[True, False]
                ),
                use_container_width=True,
                hide_index=True,
            )

        with grc_tab2:
            st.markdown("##### Control Gap Summary")
            st.dataframe(
                control_gap_view,
                use_container_width=True,
                hide_index=True,
            )

        with grc_tab3:
            st.markdown("##### Findings Requiring Manual Review")
            if unmapped_findings.empty:
                st.success("All visible findings were mapped to at least one built-in control.")
            else:
                st.dataframe(
                    unmapped_findings,
                    use_container_width=True,
                    hide_index=True,
                )

        summary = f"""
Total Findings: {len(filtered_df)}
Passed: {total_passed}
Failed: {total_failed}
Compliance: {compliance}%
Critical Issues: {critical_count}
Highest Risk Asset: {highest_risk_asset}
CIS/NIST Mapped Findings: {int((grc_df["Framework"] != "Unmapped").sum())}
"""

        st.markdown("##### Executive Summary")
        st.text_area("Executive Summary", summary, height=160, label_visibility="collapsed")

        pdf_bytes = generate_pdf(summary)
        st.download_button(
            "Download Executive Summary PDF",
            data=pdf_bytes,
            file_name="audit_report.pdf",
            mime="application/pdf",
        )
